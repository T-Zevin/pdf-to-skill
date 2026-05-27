#!/usr/bin/env python3
"""Extract text and coarse IMRaD sections for pdf-to-skill."""

from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET


WORKDIR = Path(os.environ.get("PDF_TO_SKILL_WORKDIR", Path(tempfile.gettempdir()) / "pdf_to_skill_work"))
FULL_TEXT = WORKDIR / "full_text.txt"
METADATA = WORKDIR / "metadata.json"
SECTIONS = WORKDIR / "sections.json"

SECTION_NAMES = [
    "abstract",
    "introduction",
    "background",
    "materials and methods",
    "methods",
    "method",
    "results",
    "discussion",
    "conclusion",
    "conclusions",
    "limitations",
    "data availability",
    "references",
]


def run_pdftotext(path: Path) -> str | None:
    if not shutil.which("pdftotext"):
        return None
    try:
        result = subprocess.run(
            ["pdftotext", "-layout", str(path), "-"],
            capture_output=True,
            text=True,
            timeout=180,
            check=False,
        )
    except Exception:
        return None
    return result.stdout if result.returncode == 0 and result.stdout.strip() else None


def run_pypdf(path: Path) -> str | None:
    try:
        import PyPDF2
    except Exception:
        return None

    parts: list[str] = []
    try:
        with path.open("rb") as handle:
            reader = PyPDF2.PdfReader(handle)
            for page in reader.pages:
                parts.append(page.extract_text() or "")
    except Exception:
        return None
    text = "\n".join(parts)
    return text if text.strip() else None


def run_pdfminer(path: Path) -> str | None:
    try:
        from pdfminer.high_level import extract_text
    except Exception:
        return None
    try:
        text = extract_text(str(path))
    except Exception:
        return None
    return text if text.strip() else None


def extract_pdf(path: Path) -> tuple[str, str]:
    for method, fn in (
        ("pdftotext", run_pdftotext),
        ("PyPDF2", run_pypdf),
        ("pdfminer.six", run_pdfminer),
    ):
        text = fn(path)
        if text:
            return text, method
    raise SystemExit("ERROR: could not extract text from PDF. Install poppler pdftotext, PyPDF2, or pdfminer.six.")


def read_text(path: Path) -> tuple[str, str]:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding), f"text-{encoding}"
        except UnicodeDecodeError:
            continue
    raise SystemExit(f"ERROR: could not decode text file: {path}")


def extract_docx(path: Path) -> tuple[str, str]:
    try:
        import docx

        doc = docx.Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        text = "\n".join(parts)
        if text.strip():
            return text, "python-docx"
    except Exception:
        pass

    try:
        with zipfile.ZipFile(path) as zf:
            xml = zf.read("word/document.xml")
        root = ET.fromstring(xml)
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        parts = []
        for paragraph in root.iter(f"{namespace}p"):
            texts = [node.text for node in paragraph.iter(f"{namespace}t") if node.text]
            if texts:
                parts.append("".join(texts))
        text = "\n".join(parts)
        if text.strip():
            return text, "zipfile-docx"
    except Exception:
        pass
    raise SystemExit("ERROR: could not extract text from DOCX.")


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip() + "\n"


def detect_title(text: str, filename: str) -> str:
    candidates = []
    for raw in text.splitlines()[:80]:
        line = raw.strip()
        if 12 <= len(line) <= 220 and not re.match(r"^(abstract|keywords|introduction|www\.|doi:)", line, re.I):
            if not re.search(r"^(vol\.|no\.|page|copyright|received|accepted)", line, re.I):
                candidates.append(line)
    return candidates[0] if candidates else Path(filename).stem


def detect_doi(text: str) -> str | None:
    match = re.search(r"\b10\.\d{4,9}/[-._;()/:A-Z0-9]+\b", text, re.I)
    return match.group(0).rstrip(".,;)") if match else None


def detect_sections(text: str) -> list[dict[str, object]]:
    lines = text.splitlines()
    headings: list[tuple[str, int]] = []
    names = "|".join(re.escape(name) for name in SECTION_NAMES)
    pattern = re.compile(rf"^\s*(?:\d+\.?\s*)?({names})\s*$", re.I)
    offset = 0
    for line in lines:
        stripped = line.strip()
        match = pattern.match(stripped)
        if match:
            name = match.group(1).lower()
            headings.append((name, offset))
        offset += len(line) + 1

    deduped: list[tuple[str, int]] = []
    seen_positions: set[tuple[str, int]] = set()
    for name, pos in headings:
        key = (name, pos)
        if key not in seen_positions:
            deduped.append((name, pos))
            seen_positions.add(key)

    sections = []
    for idx, (name, start) in enumerate(deduped):
        end = deduped[idx + 1][1] if idx + 1 < len(deduped) else len(text)
        sections.append({"name": name, "start": start, "end": end, "chars": end - start})
    return sections


def count_pdf_pages(path: Path) -> int | None:
    if shutil.which("pdfinfo"):
        try:
            result = subprocess.run(["pdfinfo", str(path)], capture_output=True, text=True, timeout=20, check=False)
            match = re.search(r"^Pages:\s+(\d+)", result.stdout, re.M)
            if match:
                return int(match.group(1))
        except Exception:
            pass
    try:
        import PyPDF2

        with path.open("rb") as handle:
            return len(PyPDF2.PdfReader(handle).pages)
    except Exception:
        return None


def main() -> None:
    if len(sys.argv) < 2:
        raise SystemExit("Usage: extract_paper.py <paper.pdf|paper.txt|paper.md|paper.docx>")

    source = Path(sys.argv[1]).expanduser().resolve()
    if not source.exists():
        raise SystemExit(f"ERROR: file not found: {source}")

    ext = source.suffix.lower()
    if ext == ".pdf":
        text, method = extract_pdf(source)
        pages = count_pdf_pages(source)
    elif ext in {".txt", ".md", ".markdown"}:
        text, method = read_text(source)
        pages = None
    elif ext == ".docx":
        text, method = extract_docx(source)
        pages = None
    else:
        raise SystemExit("ERROR: supported formats are .pdf, .txt, .md, .markdown, .docx")

    text = normalize_text(text)
    sections = detect_sections(text)
    words = len(text.split())
    tokens = int(words / 0.75)

    WORKDIR.mkdir(parents=True, exist_ok=True)
    FULL_TEXT.write_text(text, encoding="utf-8")
    SECTIONS.write_text(json.dumps(sections, indent=2, ensure_ascii=False), encoding="utf-8")
    METADATA.write_text(
        json.dumps(
            {
                "source_file": str(source),
                "filename": source.name,
                "format": ext.lstrip("."),
                "extraction_method": method,
                "title_guess": detect_title(text, source.name),
                "doi_guess": detect_doi(text),
                "pages": pages,
                "chars": len(text),
                "words": words,
                "estimated_tokens": tokens,
                "sections_detected": [section["name"] for section in sections],
                "output_text": str(FULL_TEXT),
                "sections_json": str(SECTIONS),
            },
            indent=2,
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    print("Extraction complete")
    print(f"  Source  : {source}")
    print(f"  Method  : {method}")
    if pages:
        print(f"  Pages   : {pages}")
    print(f"  Words   : {words:,}")
    print(f"  Tokens  : ~{tokens // 1000}K")
    print(f"  Sections: {', '.join(section['name'] for section in sections) or 'not detected'}")
    print(f"  Text    : {FULL_TEXT}")
    print(f"  Meta    : {METADATA}")
    print(f"  Sections: {SECTIONS}")


if __name__ == "__main__":
    main()
